from tkinter import *
from tkinter import messagebox
import os
import docx
import pandas as pd
from tkinter.filedialog import askdirectory
from tkinter import filedialog
from docxcompose.composer import Composer
# from win32com import client as wc





def gui():
    def selectPath():
        path_ = askdirectory()  # 使用askdirectory()方法返回文件夹的路径
        if path_ == "":
            path1.get()  # 当打开文件路径选择框后点击"取消" 输入框会清空路径，所以使用get()方法再获取一次路径
        else:
             # 实际在代码中执行的路径为“\“ 所以替换一下
            path1.set(path_)

    def select_file():
        # 单个文件选择
        selected_file_path = filedialog.askopenfilename()  # 使用askopenfilename函数选择单个文件
        file.set(selected_file_path)

    # 第1步，实例化object，建立窗口window
    window=Tk()

    # 第2步，给窗口的可视化起名字
    window.title('合并干部任免审批表')

    window.geometry('800x200')  # 这里的乘是小x

    # 桌面
    desktop=os.path.join(os.path.expanduser("~"), 'Desktop')

    strvar=StringVar()
    strvar.set(desktop)
    path1 = StringVar()
    file=StringVar()



    l1=Label(window, text="读取干部任免审批表的目录")
    e1=Entry(window, show=None, width=50, textvariable=path1)
    Button(window, text="目录选择", command=selectPath).grid(row=1, column=3)

    l2=Label(window, text="读取的名单顺序excel文件")
    e2=Entry(window, show=None, width=50, textvariable=file)
    Button(window, text="文件选择", command=select_file).grid(row=2, column=3)

    l3 = Label(window, text="合并后文件的存放路径,默认桌面")
    e3 = Entry(window, show=None, width=50, textvariable=strvar)

    l4=Label(window,text="始于全栈，成于融合，久于共创",font=("微软雅黑",8,"italic"))
    l5=Label(window,text="史祎明/技术三部/业务研发中心/ICBC",font=("微软雅黑",8,"italic"))




    def insert_point():  # 在鼠标焦点处插入输入内容
        # 读取的干部任免审批表目录
        e1value=e1.get()
        # 读取的名单顺序excel文件
        e2value=e2.get()
        # 合并后文件的存放路径
        e3value = e3.get()
        # t.insert('insert', var)

        aggExcel(e1value, e2value, e3value)





    # 第6步，创建并放置两个按钮分别触发两种情况
    btn1=Button(window, text='开始合并', width=10,
                height=2, command=insert_point)




    t=Text(window, height=3)
    l1.grid(row=1, column=1, sticky=E)
    e1.grid(row=1, column=2, sticky=E)
    l2.grid(row=2, column=1, sticky=E)
    e2.grid(row=2, column=2, sticky=E)
    l3.grid(row=3, column=1, sticky=E)
    e3.grid(row=3, column=2, sticky=E)
    btn1.grid(row=4, column=1, sticky=E)
    l4.grid(row=8, column=2, sticky=E)
    l5.grid(row=9, column=2, sticky=E)

    # 第8步，主窗口循环显示
    window.mainloop()



def aggExcel(e1, e2, e3):
    df = pd.read_excel(e2, sheet_name=0, usecols=[2], header=None)
    mylist = df.values.tolist()
    users = []
    for i in mylist:
        try:
            x = i[0].replace(" ","")
            y = findAllFile(e1, x)
            users.append(e1 + "/" + y)
        except:
            continue


    # print(users)

    # 转docx
    # for x in users:
    #     new_file = x.replace('doc','docx')
    #     new_file = new_file.replace('docm','docx')
    #     word = wc.DispatchEx('Word.Application')
    #     wc.Visible = False
    #     wc.DisplayAlerts = 0
    #     doc = word.Documents.Open(x)
    #     doc.SaveAs(new_file, 12, False, "",True, "", False, False, False,False)
    #     doc.Close()
    #     word.Quit()

    numbers_of_sections = len(users)
    master = docx.Document()
    composer = Composer(master)
    for i in users:
        doc_temp = docx.Document(i)
        doc_temp.add_page_break()
        composer.append(doc_temp)
    composer.save(os.path.join(e3,'合并后的文件.docx'))

    # new_doc = docx.Document()
    #
    # for file in users:
    #     try:
    #         print(docx.Document(file))
    #         adoc = docx.Document(file)
    #         adoc.add_page_break()
    #         for word_body in adoc.element.body:
    #             new_doc.element.body.append(word_body)
    #     except:
    #         continue
    #
    # new_doc.save(e3 + "/合并后的文件.doc")




#查找目录下是否存在某文件
def findAllFile(path, file):
    # os.listdir()方法获取文件夹名字，返回数组
    file_name_list = os.listdir(path)
    for i in file_name_list:
        if file in i:
            print(i)
            return i




if __name__ == '__main__':
    gui()

